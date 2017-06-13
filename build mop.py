from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_info import *
from docx.shared import Inches, Pt
import os
import os
import re
import socket
import sys
import netmiko
import time
from getpass import getpass

#Figure out what the interface range command is (Ex: int ra gi0/1-6) for far fewer lines put into the CLI
def get_interface_range(interfaces):
	ranges = []
	int_range = []
	for interface in interfaces:
		tmp = interface.split('net')[1]
		ranges.append(tmp)
	#print ranges
	start_range = []
	end_range = []
	for each in ranges:
		port = each.split('/')[1]
		blade = each.split('/')[0]
		next_port = int(int(port) + 1)
		next_port = str(blade) +str('/') + str(next_port)
		if next_port not in ranges:
			#print next_port
			end_range.append(each)
		last_port = int(int(port) - 1)
		
		last_port = str(blade) +str('/') + str(last_port)
		
		if last_port not in ranges:
			#start_range = each
			start_range.append(each)
			
	spot = 0
	for each in start_range:
		#print each + " - " + end_range[spot].split('/')[1]
		#spot = spot+1
		for interface in interfaces:
			if each in interface:
				int_type =  interface.split('net')[0]
		command = "interface range " + int_type + " " + each   + " - " + end_range[spot].split('/')[1]
		int_range.append(command)
		spot = spot+1
	return int_range

#Find all IPs in a string and return in a list	
def get_ip (input):
	return(re.findall(r'(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)', input))

#Get rid of the line return at the end of a line
def remove_return(entry):
	return entry.rstrip('\n')

#Read in the IP addresses we need to work on	
def read_doc (file_name):
	for line in open(file_name, 'r').readlines():
		line = remove_return(line)
		print (line)
		ips.append(line)

#Append info to a document		
def to_doc(file_name, varable):
	f=open(file_name, 'a')
	f.write(varable)
	f.close()

#Old code I don't think is used
def find_ranges(ports):
	last_port = 100
	start_port = []
	end_port = []
	for port in ports:
		tmp = port.split("/")
		tmp = tmp[1]
		tmp = int(tmp)
		if int(tmp)-1 != last_port:
			start_port.append(port)
		last_port=tmp

#Send info to the text doc		
def make_mop(hostname,ip,voice_vlan,data_vlan,final_int_range):
	name = ip+" mop info.txt"
	to_doc(name,hostname)
	to_doc(name,global_config)
	to_doc(name,add_data_helper)
	to_doc(name,add_voice_helper)
	to_doc(name,snooping_setup)
	for each in final_int_range:
		to_doc(name,each)
		to_doc(name,final_per_port)
	to_doc(name,'end \n')
	to_doc(name,'wr \n')






		
ips=[]
read_doc('IPs.txt')
print (ips)
your_user_name = input("Username: ")
your_password = getpass()	

for ip in ips:
	nac_this = []
	range_commands = []
	print(ip)
	net_connect = netmiko.ConnectHandler(device_type='cisco_ios', ip=ip, username=your_user_name, password=your_password) 
	output = net_connect.send_command_expect('show run | s interface')
	hostname = net_connect.send_command_expect('show run | i hostname')
	interfaces = output.split("interface")
	for interface in interfaces:
		if 'switchport mode access' in interface:
			port = interface.split('\n')
			nac_this.append(port[0])
			#print (port)
		if 'switchport access vlan' in interface:
			#print (interface)
			start_of_vlan = interface.find("switchport access vlan")+22
			end_of_vlan = interface[start_of_vlan:].find("\n")
			data_vlan = interface[start_of_vlan:start_of_vlan+end_of_vlan]
			#print (interface)
		if 'switchport voice vlan' in interface:
			start_of_voice_vlan = interface.find("switchport voice vlan")+22
			end_of_voice_vlan = interface[start_of_voice_vlan:].find("\n")
			voice_vlan = interface[start_of_voice_vlan:start_of_voice_vlan+end_of_voice_vlan]
	from cli_commands import *
	final_per_port = per_port_part1+data_vlan+ per_port_part2
	add_data_helper = 'interface Vlan '+ data_vlan+ helper_addresses_end
	add_voice_helper = 'interface Vlan '+ voice_vlan+ helper_addresses_end
	snooping_setup = 'ip dhcp snooping vlan ' + data_vlan+ ","+voice_vlan+snooping_config_end
	print (voice_vlan)
	print (data_vlan)
	print (ip)
	final_int_range = get_interface_range(nac_this)
	
	make_mop(hostname,ip,voice_vlan,data_vlan,final_int_range)

############make word doc############

#Add to a numbered list
def numbered_list(item):
	document.add_paragraph(item, style='List Number')

#Find the hostname of the switch
def find_hostname(file_name):
	for line in open(file_name, 'r').readlines():
		line = remove_return(line)
		if 'hostname ' in line:
			return  line.lstrip('hostname ')

#Pull in the txt file with all the CLI commands			
def read_doc (file_name):
	for line in open(file_name, 'r').readlines():
		line = remove_return(line)
		if 'hostname ' in line:
			pass
		else:
			commands.append(line)
#Get rid of the line return at the end of a line (dup code can be deleted)
def remove_return(entry):
	return entry.rstrip('\n')

#Get rid of the tabs		
def remove_tab(entry):
	return entry.lstrip('\t')

#Pulls in info from a CSV and puts it into a bullet list: If data is on the 1st cell it has no indent
#If data is in the 2ed cell it gets 1 indent, If data is on the 3ed cell it gets 2 indents
def bullet_list(csv_file):
	bul_lines = []
	for line in open(csv_file, 'r').readlines():
		if "###host_name###" in line:
			line=put_in_hostname(line,hostname)
		line = remove_return(line)
		line = line.split(',')
		bul_lines.append(line)
	for bul_line in bul_lines:
		paragraph =document.add_paragraph(bul_line[0], style='List Bullet')
		paragraph_format = paragraph.paragraph_format
		paragraph_format.space_before =  Pt(0)
		paragraph_format.space_after = Pt(0)
		try:
			if bul_line[1]:
				paragraph =document.add_paragraph(bul_line[1], style='List Bullet 2')
				paragraph_format = paragraph.paragraph_format
				paragraph_format.space_before =  Pt(0)
				paragraph_format.space_after = Pt(0)
		except:
			pass
		try:	
			if bul_line[2]:
				paragraph =document.add_paragraph(bul_line[2], style='List Bullet 3')
				paragraph_format = paragraph.paragraph_format
				paragraph_format.space_before =  Pt(0)
				paragraph_format.space_after = Pt(0)
		except:
			pass

#Replaces the place holder "###host_name###" with the actual hostname
def put_in_hostname(string,hostname):
	print(hostname)
	return string.replace("###host_name###", hostname)

#Finds all file names with that text in them that are in this folder	
def pull_file_names_with_text(text):
	file_list = []
	files = os.listdir()
	for file in files:
		if text in file:
			file_list.append(file)
	return (file_list)	
	
file_list=pull_file_names_with_text('mop info.txt')

#Goes though each major set of code and swaps out the hostname
for file in file_list:
	hostname = find_hostname(file)
	document = Document()
	commands = []
	read_doc (file)
	print (hostname)
	if "###host_name###" in head:
		head=put_in_hostname(head,hostname)
	if "###host_name###" in first_line:
		first_line=put_in_hostname(first_line,hostname)
	if "###host_name###" in second_line:
		second_line=put_in_hostname(second_line,hostname)	
	if "###host_name###" in span_info:
		span_info=put_in_hostname(span_info,hostname)	
		
		
	document.add_paragraph(head)
	#document.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	document.add_paragraph(first_line)
	
	#for item in first_builited_list:
	#	bullet_list(item)
	
	bullet_list('Participants.csv')
	
	document.add_paragraph(second_line)
	
	bullet_list('Objectives.csv')
	#for item in second_bulited_line:
	#	bullet_list(item)
	document.add_paragraph(third_line)
	
	for item in numbered_line_part_one:
		if "###host_name###" in item:
			item=put_in_hostname(item,hostname)
		numbered_list(item)
	
	document.add_paragraph(span_info)
	
	for item in numbered_line_part_two:
		if "###host_name###" in item:
			item=put_in_hostname(item,hostname)
		numbered_list(item)
	
	
	
	for item in commands:
		paragraph =document.add_paragraph(item)
		paragraph_format = paragraph.paragraph_format
		paragraph_format.space_before =  Pt(0)
		paragraph_format.space_after = Pt(0)
		
	
	
	for item in numbered_line_part_three:
		numbered_list(item)
		
	doc_to_save = file[:-4]+".docx"
	document.save(doc_to_save)
			
	